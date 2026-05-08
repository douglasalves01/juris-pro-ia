"""Testes de integração para extração e pré-processamento (sem mocks)."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from fpdf import FPDF

from api.ml.preprocessor import TextPreprocessor
from api.ml.text_extractor import TextExtractor


def _make_native_text_pdf(directory: Path) -> Path:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=12)
    bloco = (
        "Peticao inicial. DIREITO CIVIL. Contrato de prestacao de servicos. "
        "Requerente pede indenizacao por danos materiais no valor de "
        "R$ 10.000,00 conforme clausula terceira. "
        "Termos em que pede deferimento. "
    )
    texto = (bloco * 3).strip()
    pdf.multi_cell(0, 8, texto)
    path = directory / "documento_texto.pdf"
    pdf.output(str(path))
    return path


def _make_simple_docx(directory: Path) -> Path:
    path = directory / "contrato.docx"
    doc = Document()
    doc.add_heading("Contrato de honorários", level=1)
    doc.add_paragraph(
        "CONTRATANTE: Fulano de Tal, brasileiro. "
        "CONTRATADA: Escritório XYZ Advogados Associados."
    )
    doc.add_paragraph(
        "CLÁUSULA PRIMEIRA: O presente contrato regula os honorários "
        "advocatícios no valor de R$ 5.000,00."
    )
    doc.save(str(path))
    return path


def test_pdf_texto_nativo_extrai_conteudo() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = _make_native_text_pdf(Path(tmp))
        extractor = TextExtractor()
        texto = extractor.extract(str(pdf_path))
        assert "Peticao inicial" in texto
        assert "R$ 10.000,00" in texto
        meta = extractor.last_metadata
        assert meta is not None
        assert meta.file_type == "pdf"
        assert meta.has_ocr is False
        assert meta.num_pages >= 1
        assert meta.word_count > 10


def test_docx_simples_extrai_paragrafos() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = _make_simple_docx(Path(tmp))
        extractor = TextExtractor()
        texto = extractor.extract(str(docx_path))
        assert "CONTRATANTE" in texto
        assert "Escritório XYZ" in texto
        assert "R$ 5.000,00" in texto
        meta = extractor.last_metadata
        assert meta is not None
        assert meta.file_type == "docx"
        assert meta.has_ocr is False


def test_clean_preserva_valores_monetarios() -> None:
    pre = TextPreprocessor()
    bruto = (
        "O   valor   de R$ 1.234,56   será pago em 2x. "
        "Também consta R$ 99,00 no rodapé.\n\n"
        "Rodapé repetido\nRodapé repetido\nRodapé repetido\n"
    )
    limpo = pre.clean(bruto)
    assert "R$ 1.234,56" in limpo
    assert "R$ 99,00" in limpo
    assert "  " not in limpo.replace("\n\n", " ")


def test_chunking_overlap_entre_chunks() -> None:
    pre = TextPreprocessor()
    paragrafos = []
    for i in range(40):
        paragrafos.append(f"Parágrafo {i} com texto jurídico suficiente para tokenização.")
    texto = "\n\n".join(paragrafos)
    chunks = pre.split_into_chunks(texto, max_tokens=80, overlap=15)
    assert len(chunks) >= 2
    w0 = chunks[0].split()
    w1 = chunks[1].split()
    tail = w0[-15:]
    head = w1[:15]
    assert tail == head


def test_get_representative_chunk_limita_tokens() -> None:
    pre = TextPreprocessor()
    base = "palavra " * 600
    chunks = pre.split_into_chunks(base, max_tokens=200, overlap=20)
    rep = pre.get_representative_chunk(chunks, max_tokens=100)
    assert len(rep.split()) <= 100

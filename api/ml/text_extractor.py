"""Extração de texto de PDF (nativo/OCR) e DOCX com detecção de idioma."""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Final

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from pytesseract import TesseractNotFoundError
from langdetect import DetectorFactory, LangDetectException, detect_langs
from pdf2image import convert_from_path
from PIL import Image

logger = logging.getLogger(__name__)

DetectorFactory.seed = 0

TOKEN_THRESHOLD_OCR: Final[int] = 50
_MIN_CHARS_FOR_LANG_DETECT: Final[int] = 50


@dataclass(frozen=True)
class DocumentMetadata:
    """Metadados da última extração (atualizado a cada chamada a `extract`)."""

    num_pages: int
    has_ocr: bool
    file_type: str
    word_count: int
    language_code: str
    language_alert: str | None = None
    extra: dict[str, object] = field(default_factory=dict)


class TextExtractor:
    """Extrai texto de PDF (texto nativo ou OCR) e DOCX."""

    def __init__(self) -> None:
        self._last_metadata: DocumentMetadata | None = None

    @property
    def last_metadata(self) -> DocumentMetadata | None:
        return self._last_metadata

    def extract(self, file_path: str) -> str:
        path = Path(file_path).expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(f"Arquivo não encontrado: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text, meta = self._extract_pdf(path)
        elif suffix == ".docx":
            text, meta = self._extract_docx(path)
        elif suffix in (".txt", ".text"):
            text = path.read_text(encoding="utf-8", errors="ignore")
            meta = {"num_pages": 1, "has_ocr": False, "file_type": "txt"}
        else:
            raise ValueError(
                f"Formato não suportado: {suffix}. Use .pdf, .docx, .txt ou .text."
            )

        lang_code, alert = self._detect_language_info(text)
        self._last_metadata = DocumentMetadata(
            num_pages=meta["num_pages"],
            has_ocr=meta["has_ocr"],
            file_type=meta["file_type"],
            word_count=self._word_count(text),
            language_code=lang_code,
            language_alert=alert,
            extra=meta.get("extra", {}),
        )
        return text

    def _word_count(self, text: str) -> int:
        return len(text.split())

    def _count_tokens(self, text: str) -> int:
        """Aproximação de tokens por palavras (alinhada ao pipeline sem tokenizer)."""
        if not text or not text.strip():
            return 0
        return len(text.split())

    def _detect_language_info(self, text: str) -> tuple[str, str | None]:
        sample = text.strip()
        if len(sample) < _MIN_CHARS_FOR_LANG_DETECT:
            return "unknown", None
        try:
            langs = detect_langs(sample[:8000])
        except LangDetectException:
            return "unknown", None
        if not langs:
            return "unknown", None
        top = langs[0]
        code = top.lang.lower()
        prob = top.prob
        alert: str | None = None
        if code != "pt":
            alert = (
                f"Idioma detectado como '{code}' (prob. {prob:.2f}); "
                "esperado português (pt-BR) para melhor desempenho dos modelos."
            )
            logger.warning("Documento pode não estar em PT-BR: %s", alert)
        return code, alert

    def _extract_pdf(self, path: Path) -> tuple[str, dict[str, object]]:
        pages_text: list[str] = []
        num_pages = 0
        with pdfplumber.open(path) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text() or ""
                tables = [
                    md
                    for table in (page.extract_tables() or [])
                    if (md := self._table_to_markdown(table))
                ]
                if tables:
                    table_text = "\n\n[TABELA]\n".join(tables)
                    t = (
                        f"{t}\n\n[TABELA]\n{table_text}"
                        if t.strip()
                        else f"[TABELA]\n{table_text}"
                    )
                pages_text.append(t)
        raw = "\n\n".join(pages_text)
        has_ocr = False
        if self._count_tokens(raw) < TOKEN_THRESHOLD_OCR:
            ocr_text, ocr_pages = self._ocr_pdf(path)
            if ocr_text.strip():
                raw = ocr_text
                has_ocr = True
                num_pages = ocr_pages
        return raw, {
            "num_pages": num_pages,
            "has_ocr": has_ocr,
            "file_type": "pdf",
            "extra": {},
        }

    def _ocr_pdf(self, path: Path) -> tuple[str, int]:
        try:
            images = convert_from_path(
                str(path),
                dpi=200,
                fmt="png",
            )
        except Exception as exc:
            logger.error("Falha ao converter PDF em imagens: %s", exc)
            return "", 0
        parts: list[str] = []
        for img in images:
            if not isinstance(img, Image.Image):
                continue
            gray = img.convert("L")
            # Melhora contraste e nitidez para OCR mais preciso
            from PIL import ImageOps, ImageFilter
            gray = ImageOps.autocontrast(gray, cutoff=2)
            gray = gray.filter(ImageFilter.SHARPEN)
            try:
                txt = pytesseract.image_to_string(
                    gray, lang="por+eng",
                    config="--psm 6 --oem 3",
                )
            except TesseractNotFoundError:
                logger.warning(
                    "Tesseract não encontrado no PATH; OCR indisponível."
                )
                return "", 0
            except Exception as exc:
                logger.warning("OCR falhou em uma página do PDF: %s", exc)
                txt = ""
            parts.append(txt)
        return "\n\n".join(parts), len(images)

    def _table_to_markdown(self, table: list[list[object | None]]) -> str:
        rows: list[list[str]] = []
        for row in table or []:
            cells = [
                str(cell).strip()
                for cell in (row or [])
                if cell is not None and str(cell).strip()
            ]
            if len(cells) >= 2:
                rows.append(cells)

        if not rows:
            return ""

        width = max(len(row) for row in rows)
        if width < 2:
            return ""

        normalized = [row + [""] * (width - len(row)) for row in rows]
        header = normalized[0]
        separator = ["-" * max(3, len(cell)) for cell in header]
        body = normalized[1:]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        lines.extend("| " + " | ".join(row) + " |" for row in body)
        return "\n".join(lines)

    def _extract_docx(self, path: Path) -> tuple[str, dict[str, object]]:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    paragraphs.append(line)
        text = "\n\n".join(paragraphs)
        return text, {
            "num_pages": self._estimate_docx_pages(text),
            "has_ocr": False,
            "file_type": "docx",
            "extra": {},
        }

    def _estimate_docx_pages(self, text: str) -> int:
        """Estimativa quando não há paginação nativa no DOCX."""
        words = self._word_count(text)
        if words == 0:
            return 1
        estimated = max(1, round(words / 300))
        return estimated
